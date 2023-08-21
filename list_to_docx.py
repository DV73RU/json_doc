from docx import Document
from htmldocx import HtmlToDocx
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import html2text

# Пример HTML-кода для вставки
# html_text = "<p>This is <strong>bold</strong> text.</p> <a href=\"https://news.tpprf.ru/ru/news/4082526/\">6 марта 2023 года</a>"
#
# # Создание объекта Document
# doc = Document()
#
# # Инициализация конвертера
# html_to_docx = HtmlToDocx()
#
# # Преобразование HTML в DOCX и вставка в документ
# html_to_docx.add_html_to_document(html_text, doc)
#
# # Сохранение документа
# doc.save("output.docx")

from docx import Document
from htmldocx import HtmlToDocx
from bs4 import BeautifulSoup
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Пример JSON-блока
block = {
    "blockType": 11,
    "elemList": {
        "elemType": 2,
        "elems": [
            "<strong>'Элемент списка 1'</strong> Текст 1",
            "<strong>'Элемент списка 2'</strong> Текст 2",
            "<strong>'Элемент списка 3'</strong> Текст 3",
            "<strong>'Элемент списка 4'</strong> Текст 4"
        ]
    }
}

# Создание объекта Document
doc = Document()

# Создание экземпляра конвертера из HTML в DOCX
html_to_docx = HtmlToDocx()

# Извлечение элементов списка из JSON-блока
list_items_html = block["elemList"]["elems"]

# Добавление элементов списка в документ DOCX
for item_html in list_items_html:
    # Создаем параграф для элемента списка
    paragraph = doc.add_paragraph(style="List Bullet")

    # # Добавляем знак "-" перед тегом <strong>
    # run = paragraph.add_run("- ")
    # run.bold = False

    # Преобразуем HTML в текст и добавляем его в параграф
    soup = BeautifulSoup(item_html, "html.parser")
    text = soup.get_text()
    paragraph.add_run(text)

    # Устанавливаем стиль для текста в параграфе
    for run in paragraph.runs:
        run.font.size = Pt(12)
        run.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Сохранение документа
doc.save("output.docx")


