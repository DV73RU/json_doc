import os
import hashlib
import requests
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from tqdm import tqdm
from htmldocx import HtmlToDocx
from docx import Document
from docx.shared import Pt
import html2text
from read_list_json import read_links_from_file
import logging

"""
Создаем директорию с названием новости
Создаем документ *.docx с названием новости
Скачиваем картинки из новости и помещаем в директорию названия новости
Скачиваем доп материалы к новости и сохраняем в директорию 'название новости/matetials'

"""
file_path = "list_json.txt"
# Получаем список ссылок из файла
links = read_links_from_file(file_path)
# Логирование ошибок
logging.basicConfig(filename='err_log.log', level=logging.ERROR, format='%(asctime)s - %(levelname)s - %(message)s',
                    encoding='utf-8')


# Список URL для JSON данных
json_urls = links


def add_empty_line(doc):  # Пустая строка
    doc.add_paragraph("")


def html_to_plain_text(html):
    h = html2text.HTML2Text()
    h.ignore_links = True
    return h.handle(html)


def add_text_block(doc, text, font_size, font_style=None, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT,
                   spacing_after=Pt(0)):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    if font_style:
        if "bold" in font_style:
            run.bold = True
        if "italic" in font_style:
            run.italic = True

    run.font.size = Pt(font_size)
    run.text = text

    paragraph.alignment = alignment
    paragraph.paragraph_format.space_after = spacing_after
    return paragraph


def clean_filename(filename):
    invalid_chars = r'\/:*?"<>|'  # Здесь перечислены недопустимые символы в названии директории и файла
    cleaned_filename = ''.join(c for c in filename if c not in invalid_chars)
    return cleaned_filename


# Для каждого URL получаем JSON данные и создаем документ

# Статус-бар для JSON-файлов
total_json_urls = len(json_urls)
json_bar = tqdm(total=total_json_urls, desc='Процесс JSON URLs')
for json_url in json_urls:
    try:
        response = requests.get(json_url)
        response.raise_for_status()  # Проверка на успешный ответ (код 200)
        your_json = response.json()
        json_bar.update(1)

    except requests.exceptions.RequestException as e:
        print(f"Ошибка при получении данных из URL: {json_url}")
        print(f"Ошибка: {e}")
        logging.error(f"Ошибка при получении данных из URL: {json_url}")
        continue  # Прерываем итерацию и переходим к следующему URL

    material_folder_name = "Материалы"

    # Создание документа
    doc = Document()
    # Создание объекта конвертера
    html_to_docx = HtmlToDocx()

    # Получение значения из поля "id"
    id_ = your_json["data"]["id"]
    add_text_block(doc, str(id_), 10, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)  # Добавление названия ID статьи
    print(f"ID статьи: {id_}")

    # Создаем директорию для сохранения документов, если её еще нет
    # Название директории название слитьи и номер id
    title = your_json["data"]["title"]
    title_for_folder = clean_filename(title + "_" + str(id_))  # Используем функцию для обработки имени
    folder_path = os.path.join(os.getcwd(), material_folder_name, title_for_folder)

    os.makedirs(folder_path, exist_ok=True)
    print(f"Создана директория: {folder_path}")

    # Путь к файлу .docx внутри папки
    docx_filename = os.path.join(folder_path, f"{title_for_folder}.docx")
    print(f"Создан файл: {docx_filename}")

    doc.styles['Normal'].font.name = 'Times New Roman'
    # Добавление таблицы для вставки значений Заголовок и Дискриптион
    table = doc.add_table(rows=3, cols=2)
    table.autofit = True
    # Установка стиля границ таблицы для тонких линий
    table.style = 'Table Grid'
    # Вставка заголовков в первый столбец
    table.cell(0, 0).text = "Title"
    table.cell(1, 0).text = "Description"
    print("Добавлена таблица")

    # Добавление названия статьи
    add_text_block(doc, title, 14, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
    print(f"Title стать: {title}")

    # Получение значения из поля "subtitle"
    subtitle = your_json["data"]["subtitle"]
    print(f"Subtitle статьи: {subtitle}")

    # Получение значения из поля "seoDescription"
    seo_description = your_json["data"]["rubric"]["seoDescription"]
    # add_text_block(doc, seo_description, 12, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
    print(f"seoDescription: {seo_description}")

    add_text_block(doc, subtitle, 12, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)  # Подзаголовок
    add_empty_line(doc)  # Вызов функции для добавления пустой строки

    # Получение значения "materials"
    materials = your_json["data"]["materials"]
    print(f"Доп материалы: {materials}")

    # Вставка значений title и seo_description во второй столбец
    table.cell(0, 1).text = title
    table.cell(1, 1).text = seo_description

    table.cell(2, 0).text = "Дополнительные материалы"
    # table.cell(2, 1).text = materials_info

    # Создание папки для дополнительных материалов
    materials_folder = os.path.join(folder_path, "materials")
    os.makedirs(materials_folder, exist_ok=True)
    # Скачивание дополнительных материалов (если есть)
    if materials:
        for material in materials:
            material_name = material["name"]
            material_url = material["file"]
            material_extension = material_url.split(".")[-1]
            material_filename = f"{material_name}.{material_extension}"
            material_path = os.path.join(materials_folder, material_filename)
            try:
                material_response = requests.get(material_url)
                material_response.raise_for_status()  # Проверка на успешный ответ (код 200)
                with open(material_path, "wb") as material_file:
                    material_file.write(material_response.content)
                print(f"Дополнительный материал скачан: {material_path}")
                materials_info = material_name
                table.cell(2, 1).text = materials_info  # Добавление в таблицу инфу о доп материале
            except requests.exceptions.RequestException as e:

                error_message = f"Ошибка при скачивании материала {material_name}: {e}"
                logging.error(f"Ошибка при скачивании материала {material_name}: {e} : id = {id_}")
                colored_error_message = f"\033[91m{error_message}\033[0m"
                print(colored_error_message)
                materials_info = f"Ошибка скачивания материала: {material_name}"

                # Создаем объект run для стилизации текста
                run = table.cell(2, 1).paragraphs[0].add_run(materials_info)

                # Применяем стили шрифта (красный цвет)
                font = run.font
                font.color.rgb = RGBColor(255, 0, 0)  # Красный цвет
    else:
        print("Нет дополнительных материалов.")
        materials_info = "Нет дополнительных материалов"  # Добавление информации в таблицу о доп материалах
        table.cell(2, 1).text = materials_info

    # Добавление текстовых блоков из "blocks"
    for block in your_json["data"]["blocks"]:
        block_type = block["blockType"]
        if block_type == 1:  # Основной текст
            text = block.get("text")

            if text:
                # Преобразование HTML в DOCX
                html_to_docx.add_html_to_document(text, doc)
                print(f"Добавлен: Основной текст")


        # elif block_type == 10:  # Заголовок V1.0
        #     text = block.get("text")
        #     if text:
        #         # add_text_block(doc, text, 16, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
        #         html_to_docx.add_html_to_document(text, doc,)
        #         print(f"Добавлен: Заголовок")
        #         # add_empty_line(doc)

        elif block_type == 10:  # Заголовок V1.1
            text = block.get("text")
            if text:
                paragraph = doc.add_paragraph(text)
                run = paragraph.runs[0]
                run.font.size = Pt(13)
                run.bold = True
                print(f"Добавлен: Заголовок")


        elif block["blockType"] == 11:   # Список V.1
            # Извлекаем элементы списка
            list_items_html = block["elemList"]["elems"]
            for item in list_items_html:
                html_to_docx.add_html_to_document("• " + item, doc)
                print(f"Добавлен: Элемент списка")

        elif block["blockType"] == 24:   # Заголовок списка V.1
            text = block.get("text")
            if text:
                paragraph = doc.add_paragraph(text)
                run = paragraph.runs[0]
                run.font.size = Pt(13)
                run.bold = True
                print(f"Добавлен: Заголовок списка")

        elif block_type == 2:  # Цитата и её автор V.2
            text = block.get("text")    # Текст цитаты
            author = block.get("author")    # Автор
            comment = block.get("comment")
            regalia = block.get("regalia")  # Регалии автора

            if text:
                # Добавление текста с курсивным стилем
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(f'"{text}"')    # Добавляем кавычки текст цитаты
                font = run.font
                font.italic = True
                # doc.add_paragraph(text)
                print("Добавлен: Текст цитаты")

            if author:
                author_paragraph = doc.add_paragraph()
                author_run = author_paragraph.add_run(author)
                author_font = author_run.font
                author_font.bold = True
                print("Добавлен: Автор цитаты")

            if regalia:
                regalia_paragraph = doc.add_paragraph()
                author_run = regalia_paragraph.add_run(regalia)
                regalia_font = author_run.font
                regalia_font.bold = True
                regalia_font.italic = True  # Установка курсивного стиля
                regalia_font.size = Pt(10)  # Установка размера шрифта в 10 точек
                print("Добавлен: Регалии автора")
            #
            # if comment:
            #     doc.add_paragraph('Комментарий пользователя: "' + comment + '"')
            #     print("Добавлен: Комментарий блока 2")

        elif block["blockType"] == 5 and "carousel" in block:  # Картинки
            carousel_images = block["carousel"]

            for index, carousel_item in enumerate(carousel_images):

                image_url = carousel_item["image"]
                image_response = requests.get(image_url)

                image_extension = image_url.split(".")[-1]
                image_hash = hashlib.md5(image_response.content).hexdigest()
                image_filename = f"carousel_{image_hash}.{image_extension}"
                image_path = os.path.join(folder_path, image_filename)
                with open(image_path, "wb") as image_file:
                    image_file.write(image_response.content)
                try:
                    image_response.raise_for_status()  # Проверка на успешный ответ (код 200)
                    with open(image_path, "wb") as image_file:
                        image_file.write(image_response.content)
                    print(f"Изображение из блока 'blockType:5''carousel' сохранено: {image_path}")

                except requests.exceptions.RequestException as e:
                    print(f"Ошибка при скачивании изображения {image_url}: {e}")
                    error_message = f"Ошибка при скачивании изображения {image_url}: {e}"
                    logging.error(f"Ошибка при скачивании изображения {image_url}: {e} : id = {id_}")
                    colored_error_message = f"\033[91m{error_message}\033[0m"
                    print(colored_error_message)

                # Добавление подписи к изображению
                sign = carousel_item.get("sign")
                if sign is not None:
                    add_empty_line(doc)
                    # Вставка изображения в docx
                    doc.add_picture(image_path, width=Inches(6.0))  # Изменение размеров картинки
                    print(f"Добавлен: Изображение")
                    if sign:
                        # Добавление комментария, если он есть
                        # add_text_block(doc, "Комментарий:", 12, WD_PARAGRAPH_ALIGNMENT.LEFT)
                        add_text_block(doc, sign, 10, font_style='italic', alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
                        print(f"Добавлен: Комментарий к изображению")
                    add_empty_line(doc)
                else:
                    # Вставка изображения в docx без комментария
                    doc.add_picture(image_path, width=Inches(6.0))  # Изменение размеров картинки
                    # add_empty_line(doc)
                    print(f"Добавлен: Изображение без комментария")

        elif block["blockType"] == 3:  # Ведео


            # Скачивание и вставка изображения
            image_url = block.get("image")
            if image_url:
                response = requests.get(image_url)
                if response.status_code == 200:
                    image_filename = "screenshot_video.png"  # Имя файла изображения
                    image_path = os.path.join(folder_path, image_filename)  # Полный путь к файлу
                    print(f"Скачено: скрин видео {image_filename}")
                    with open(image_path, "wb") as f:
                        f.write(response.content)
                    doc.add_picture(image_path, width=Inches(6.0))
                    print(f"Добавлено: Скрин видео {image_filename}")

                # Вставка описания к изображению
                title = block.get("title")
                if title:
                    paragraph = doc.add_paragraph()
                    paragraph.add_run(title).font.size = Pt(12)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    print("Добавлен: Описание к видео.")

                # Вставка ссылки на видео
            link = block.get("link")
            if link:
                modified_link = link.replace("/_HLS_/", "/").replace(".m3u8", ".mp4")
                modified_link = modified_link[:modified_link.rfind("/")] + ".mp4"

                response = requests.head(modified_link)
                if response.status_code == 200:
                    doc.add_paragraph("Ссылка на видео:")
                    doc.add_paragraph(modified_link)
                    print(f"Добавлен: url видео - {modified_link}")
                else:
                    doc.add_paragraph("Видео недоступно по указанной ссылке:")
                    doc.add_paragraph(modified_link)
                    logging.error(f"Видео недоступно по ссылке: {modified_link} {id_}")
                    print("Видео недоступно по указанной ссылке.")

                # doc.add_paragraph("Исходная ссылка на видео:")
                # print(f"Исходная ссылка на видео - {link}")

    # Сохранение документа названием статьи в папку с названием статьи
    doc.save(docx_filename)
    print(
        f"Документ сохранен: {docx_filename}\n------------------------------------------------------------------------------")
    json_bar.close()
